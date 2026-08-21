# -*- coding: utf-8 -*-
"""生成 王栋 的 Word 简历（AI产品构建者方向）。

用法：python generate_resume.py
输出：王栋-简历.docx
修改内容后重跑即可重新生成。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 常量 ----------
FONT_CN = "微软雅黑"
FONT_EN = "Calibri"
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_SUB = RGBColor(0x6B, 0x65, 0x60)
COLOR_ACCENT = RGBColor(0x8B, 0x76, 0x40)   # 金棕
COLOR_LINE = RGBColor(0xC4, 0xA9, 0x75)     # 亮金
COLOR_TAG = RGBColor(0x5A, 0x50, 0x45)


def set_font(run, size=10, bold=False, color=COLOR_TEXT, italic=False, cn=FONT_CN):
    """设置中英文字体与字号。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def add_bottom_border(paragraph, color="C4A975", size="8"):
    """给段落加底部边框线。"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def para(doc, text="", size=10, bold=False, color=COLOR_TEXT, italic=False,
         align=None, space_before=0, space_after=4, line=1.3, cn=FONT_CN):
    """添加段落并返回。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color, italic=italic, cn=cn)
    return p


def section_title(doc, text):
    """章节标题：加粗 + 金色下划线。"""
    p = para(doc, text, size=12, bold=True, color=COLOR_ACCENT,
             space_before=14, space_after=8, line=1.2)
    add_bottom_border(p)
    return p


def project_block(doc, name, tagline, bullets, tags, star=False, scale_note=None):
    """项目经历块：项目名 + 一句话 + bullet + 标签行。"""
    # 项目名
    p = para(doc, space_before=8, space_after=2, line=1.2)
    title = ("★ " if star else "") + name
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=COLOR_ACCENT if star else COLOR_TEXT)
    # 一句话
    para(doc, tagline, size=9.5, color=COLOR_SUB, space_after=2, line=1.3)
    # bullets
    for b in bullets:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.3
        pf.left_indent = Cm(0.5)
        r = p.add_run("• " + b)
        set_font(r, size=9.5, color=COLOR_TEXT)
    # 标签行
    if tags:
        p = para(doc, space_before=2, space_after=2, line=1.2)
        r = p.add_run("  |  ".join(tags))
        set_font(r, size=8.5, color=COLOR_TAG, italic=True)
    if scale_note:
        para(doc, scale_note, size=8.5, color=COLOR_TAG, space_after=6, line=1.2)


def job_block(doc, period, company_role, desc_lines):
    """工作经历块。"""
    p = para(doc, space_before=8, space_after=2, line=1.2)
    r = p.add_run(period + "  ")
    set_font(r, size=9.5, bold=True, color=COLOR_ACCENT)
    r2 = p.add_run(company_role)
    set_font(r2, size=10.5, bold=True, color=COLOR_TEXT)
    for d in desc_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.3
        pf.left_indent = Cm(0.5)
        r = p.add_run("• " + d)
        set_font(r, size=9.5, color=COLOR_TEXT)


# ============================================================
doc = Document()

# 页面设置 A4
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

# ============ 头部 ============
para(doc, "王栋", size=22, bold=True, space_after=2, line=1.1)
para(doc, "AI产品构建者  ·  如何用AI去提高效率", size=11,
     color=COLOR_ACCENT, space_after=4, line=1.2)

para(doc, "📧 1252395926@qq.com   📞 130 3283 9382   🔗 github.com/wd-cloud-spec   🔗 gitee.com/scyilang_0",
     size=9, color=COLOR_SUB, space_after=2, line=1.2)
para(doc, "📍 重庆 · 成都  ·  随时到岗  ·  30岁  ·  电子科技大学（985）本科",
     size=9, color=COLOR_SUB, space_after=8, line=1.2)

# ============ 一句话定位 ============
section_title(doc, "一句话定位")
para(doc, "7年跨界：平面设计3年（审美与用户体验）→ B端产品经理1年（企业痛点与方案输出）→ 研学创业3年（从0到1的完整商业闭环，近百场交付、零安全事故）→ 独立构建多个AI产品（把大模型能力变成可交付、可变现的应用）。",
     size=10, color=COLOR_TEXT, space_after=4, line=1.35)

# ============ 核心能力（三栏表格）============
section_title(doc, "核心能力")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

caps = [
    ("商业闭环", "创业3年，独立完成产品研发、渠道拓客、商务谈判、交付回款、客户维护全链路。会算账、能控本、懂风控。"),
    ("技术落地", "独立完成AI产品全栈开发：Python/Node.js/JavaScript/PySide6，RAG选型、OCR降级链、多Agent编排均有实战。"),
    ("设计能力", "平面设计3年，兼具B端产品与售前思维，从原型到视觉落地全流程，独立输出界面与投标物料。"),
]
for i, (title, desc) in enumerate(caps):
    cell = table.cell(0, i)
    cell.width = Cm(5.8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=COLOR_ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(desc)
    set_font(r2, size=9, color=COLOR_TEXT)

# ============ AI项目经历 ============
section_title(doc, "AI 项目经历")

project_block(
    doc, "DeepSeek Harness Mobile — 手机远程指挥台（旗舰项目）",
    "AI 在电脑上干活的时候，你不用守在电脑前——手机上就能看到它做到哪一步了、需要你拍板的地方点一下就行。",
    [
        "实时看进度：AI 正在做什么、做到哪了，手机上一目了然",
        "远程发话 + 审批放行：给 AI 安排任务或补充要求，手机上直接说；需要确认的事点一下批准",
        "断点续跑：中途断线也不怕，回来后接着原来的进度继续",
        "安全设防：只有你的手机能指挥，别人偷不走控制权（14 路径白名单 + 旋转 Token 安全桥接）",
        "技术底子：双面插件架构，宿主端 + 浏览器 UI 同一包分发，~2,650 行核心代码，MIT 开源",
    ],
    ["DeepSeek Harness", "SSE", "移动端", "安全桥接", "Agent控制"],
    star=True,
    scale_note="工程规模：~2,650行 · 双面插件 · 单文件移动端UI · 12条工程踩坑记录",
)

project_block(
    doc, "CodeClean — AI 代码清理平台（旗舰项目）",
    "项目代码堆久了会有大量\"没人敢动\"的老文件。CodeClean 让不懂代码的人也能安全清理——先试跑验证、删了能找回、出问题自动还原。",
    [
        "傻瓜式清理：不用看懂代码，AI 自动找出冗余文件；清理前先在隔离环境试跑，确认不影响功能才动手",
        "后悔药：删掉的文件进\"回收站\"，一条命令随时找回；清理后自动校验，异常整体回滚",
        "AI 会学习：清理经验自动沉淀，下次更准、更少出错",
        "技术底子：12 Agent 协作（清理端 7 个 + 学习端 5 个），10 节点 LangGraph 编排，断点续跑",
        "三道防线：沙箱副本只读执行 + 只隔离不删除（SHA256 溯源可回迁）+ 基线回归自动回滚",
    ],
    ["Python", "LangGraph", "12 Agent", "Tree-sitter", "沙箱隔离"],
    star=True,
    scale_note="工程规模：5000行 Python · 36个模块 · 12个Prompt模板 · 85项自动化测试全过",
)

project_block(
    doc, "铜雀台 × 商单台 — AI 销售管理平台（旗舰项目）",
    "一套系统管完销售的\"事前事后\"：铜雀台管客户（建档、跟进、判意向），商单台管履约（报价→合同→回款→交付→售后），最方便的是：说话就能干活。",
    [
        "说话即操作：对系统说\"把高配工位套餐 2 套报价给张三\"，报价单自动生成、金额自动算好（实测自动计价 21,897 元）",
        "客户不重复录：信息录一次，两边系统自动共享；报价单生成只读链接，客户免登录查看",
        "丢单有预警：客户久无动静自动提醒；成交有审批把关，从报价到售后全流程记录",
        "技术底子：语音→动作流水线自主开发（讯飞 ASR → DeepSeek 意图 → 自动执行），金额一致性由数据库事务保证",
        "版本化幂等迁移：从 2 人团队演化到 20-50 人组织（层级权限 + 公海自动回收）",
    ],
    ["Next.js", "React", "Supabase", "DeepSeek", "讯飞ASR", "PWA"],
    star=True,
    scale_note="工程规模：两个应用 ~15,000行 · 已部署 Vercel · PWA可安装 · 真实可上线",
)

project_block(
    doc, "成绩统计工具 — AI×教育桌面应用",
    "给老师用的\"拍照录成绩\"工具：拍一张成绩单照片，AI 自动识别成电子表格，分析图表自动生成。",
    [
        "拍照即录入：不用手敲，拍个照成绩全进去了；多张照片自动按姓名合并",
        "分析自动出：平均分、排名、分数段分布等图表一键生成；错题关键词检索",
        "断网也能用：4 级 OCR 降级链兜底（Qwen-VL → 腾讯云 → 本地 OCR）",
        "数据不出校门：成绩只存在本地电脑；双击即用的 .exe，教师零学习成本",
    ],
    ["Python", "Qwen-VL", "PySide6", "Fluent Design", "SQLite"],
)

project_block(
    doc, "企业管理系统 — 中小企业AI全栈管理平台",
    "中小企业的一站式管理后台——办公（OA）、客户（CRM）、仓库（WMS）三合一，AI 把重复琐事全干了。",
    [
        "6 大 AI 功能：智能填单（说一句\"申请办公用品采购\"表单自动填好）、发票自动识别、跟进摘要、赢单预测、采购建议、全局 AI 助手",
        "一个 AI 助手：想问什么直接问，系统里所有信息它都知道",
        "技术底子：Node.js 零依赖后端，100 并发压测零崩溃；统一认证 + 三级权限（17 项细粒度权限）",
        "可局域网/私有云部署，数据留在自己手里",
    ],
    ["Node.js", "DeepSeek", "Electron", "RBAC", "零依赖"],
)

# ============ 工作经历 ============
section_title(doc, "工作经历")

job_block(
    doc, "2023.08 - 至今",
    "高校研学项目 · 创始人 / 项目负责人（自主创业）",
    [
        "独立完成产品研发→渠道拓客→商务谈判→项目交付→客户维护→复购的完整商业闭环，深耕成都高校研学市场",
        "累计落地近百场研学实践，师生满意度95%以上，零安全事故、零投诉",
        "独立核算成本利润、统筹团队排班培训，多场次项目同步有序落地",
    ],
)

job_block(
    doc, "2022.05 - 2023.08",
    "信锐科技 · B端产品经理（政企/高校园区网络）",
    [
        "对接政企与高校客户，挖掘网络建设与信息化改造需求，撰写解决方案、技术白皮书与招投标资料",
        "跟进产品迭代与项目落地，联动技术、实施团队解决执行中的对接问题",
    ],
)

job_block(
    doc, "2019.01 - 2022.05",
    "淘宝美工 / 平面设计师",
    [
        "3年全品类视觉设计：电商店铺装修、品牌画册、活动物料、招投标画册等全案输出",
        "具备独立完成产品界面设计的能力，保证项目视觉交付不依赖外部设计师",
    ],
)

# ============ 教育背景 ============
section_title(doc, "教育背景")
para(doc, "电子科技大学（985/211） · 电波传播与天线专业 · 本科  |  2014.09 - 2018.06",
     size=10, color=COLOR_TEXT, space_after=2, line=1.3)

# ============ 附加信息 ============
section_title(doc, "附加信息")
para(doc, "个人主页：wd-cloud-spec.github.io（AI项目视频演示）  ·  GitHub：github.com/wd-cloud-spec  ·  随时到岗  ·  目标城市：重庆、成都",
     size=9.5, color=COLOR_SUB, space_after=2, line=1.3)

out = "王栋-简历.docx"
doc.save(out)
print(f"已生成：{out}")
